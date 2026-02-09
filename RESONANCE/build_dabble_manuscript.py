#!/usr/bin/env python3
"""
Dabble Writer Manuscript Builder
Compiles all chapter files into a .docx with proper italics formatting.
Paste from Word/Google Docs into Dabble to preserve formatting.
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Note: python-docx not installed. Run: pip install python-docx")

CHAPTERS_DIR = Path("/workspaces/pilot/RESONANCE/chapters")
OUTPUT_DIR = Path("/workspaces/pilot/RESONANCE")

def get_chapter_sort_key(filename):
    """Extract chapter number for sorting, handling sub-chapters like CH8a, CH8b"""
    match = re.search(r'CH(\d+)([a-z])?', filename)
    if match:
        num = int(match.group(1))
        sub = match.group(2) or ''
        # Convert letter to decimal (a=0.1, b=0.2, etc.)
        sub_val = (ord(sub) - ord('a') + 1) * 0.1 if sub else 0
        return num + sub_val
    return 999  # Put unmatched files at end

def clean_chapter_content(content, filename):
    """Ensure chapter follows Dabble format rules"""
    lines = content.strip().split('\n')

    # Remove any trailing whitespace from each line
    lines = [line.rstrip() for line in lines]

    # Ensure proper ending (single blank line concept - just clean ending)
    while lines and lines[-1] == '':
        lines.pop()

    return '\n'.join(lines)

def add_formatted_paragraph(doc, text):
    """Add a paragraph with _underscore_ converted to italics"""
    para = doc.add_paragraph()

    # Pattern to find _italic text_
    pattern = r'_([^_]+)_'
    last_end = 0

    for match in re.finditer(pattern, text):
        # Add text before the match (normal)
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])

        # Add the matched text (italic)
        italic_run = para.add_run(match.group(1))
        italic_run.italic = True

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        para.add_run(text[last_end:])

def build_manuscript():
    """Build complete manuscript from chapter files"""

    if not DOCX_AVAILABLE:
        print("ERROR: python-docx is required. Install with: pip install python-docx")
        return None

    # Find all chapter files
    chapter_files = list(CHAPTERS_DIR.glob("RESONANCE_CH*.txt"))

    # Filter out scaffolds, old versions, transcripts
    chapter_files = [f for f in chapter_files if 'SCAFFOLD' not in f.name
                     and '_old' not in f.name
                     and '_TRANSCRIPT' not in f.name]

    # Sort by chapter number
    chapter_files.sort(key=lambda f: get_chapter_sort_key(f.name))

    print(f"Found {len(chapter_files)} chapters to compile:\n")

    # Create Word document
    doc = Document()
    total_words = 0

    for i, chapter_file in enumerate(chapter_files, 1):
        content = chapter_file.read_text(encoding='utf-8')
        cleaned = clean_chapter_content(content, chapter_file.name)
        word_count = len(cleaned.split())
        total_words += word_count

        print(f"  {i:2}. {chapter_file.name} ({word_count:,} words)")

        # Add each line as a paragraph with italic formatting
        lines = cleaned.split('\n')
        for line in lines:
            if line.strip():
                add_formatted_paragraph(doc, line)
            else:
                doc.add_paragraph()  # Empty paragraph for blank lines

        # Add chapter separator (Dabble recognizes ### as scene break)
        if i < len(chapter_files):
            doc.add_paragraph()
            doc.add_paragraph()

    # Generate output filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = OUTPUT_DIR / f"RESONANCE_DABBLE_{timestamp}.docx"

    # Save document
    doc.save(str(output_file))

    print(f"\n{'='*60}")
    print(f"MANUSCRIPT BUILT SUCCESSFULLY")
    print(f"{'='*60}")
    print(f"  Chapters: {len(chapter_files)}")
    print(f"  Total words: {total_words:,}")
    print(f"  Output: {output_file.name}")
    print(f"\nOpen in Word/Google Docs, then copy/paste into Dabble.")
    print(f"Italics will be preserved!")

    return output_file

if __name__ == "__main__":
    build_manuscript()
