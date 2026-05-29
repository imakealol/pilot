#!/usr/bin/env python3
"""
Builds complete manuscript including CH24a and the Epilogue.
Converts _underscore_ markers to real italics.
Paste from Word/Google Docs into Dabble to preserve formatting.
"""

import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Note: python-docx not installed. Run: pip install python-docx")

CHAPTERS_DIR = Path("/workspaces/pilot/RESONANCE/chapters")
OUTPUT_DIR = Path("/workspaces/pilot/RESONANCE")

def get_sort_key(filename):
    """Sort by chapter number. Epilogue sorts last."""
    if 'RESONANCE_EP_' in filename:
        return 9999
    match = re.search(r'CH(\d+)([a-z])?', filename)
    if match:
        num = int(match.group(1))
        sub = match.group(2) or ''
        sub_val = (ord(sub) - ord('a') + 1) * 0.1 if sub else 0
        return num + sub_val
    return 9998

def add_formatted_paragraph(doc, text):
    para = doc.add_paragraph()
    pattern = r'_([^_]+)_|\*([^*]+)\*'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])
        run = para.add_run(match.group(1) or match.group(2))
        run.italic = True
        last_end = match.end()
    if last_end < len(text):
        para.add_run(text[last_end:])

def build():
    if not DOCX_AVAILABLE:
        print("ERROR: python-docx required. Install with: pip install python-docx")
        return

    # Grab all chapter files + epilogue
    files = (
        list(CHAPTERS_DIR.glob("RESONANCE_CH*.txt")) +
        list(CHAPTERS_DIR.glob("RESONANCE_EP_*.txt"))
    )

    # Filter scaffolds and old versions
    files = [f for f in files if 'SCAFFOLD' not in f.name
             and '_old' not in f.name
             and '_TRANSCRIPT' not in f.name]

    files.sort(key=lambda f: get_sort_key(f.name))

    doc = Document()
    total_words = 0

    for i, path in enumerate(files):
        content = path.read_text(encoding='utf-8').strip()
        word_count = len(content.split())
        total_words += word_count
        print(f"  {i+1:2}. {path.name} ({word_count:,} words)")

        for line in content.split('\n'):
            if line.strip() == '---':
                continue
            elif line.strip():
                add_formatted_paragraph(doc, line)
            else:
                doc.add_paragraph()

        if i < len(files) - 1:
            doc.add_paragraph()
            doc.add_paragraph()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output = OUTPUT_DIR / f"RESONANCE_COMPLETE_{timestamp}.docx"
    doc.save(str(output))

    print(f"\nBuilt: {output.name}")
    print(f"Chapters: {len(files)} | Total words: {total_words:,}")
    print("Open in Word/Google Docs, then paste into Dabble.")

if __name__ == "__main__":
    build()
